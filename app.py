import streamlit as st
import pandas as pd
import re
from io import BytesIO
from docx import Document
from docx.shared import RGBColor, Pt
from Bio.SeqUtils.ProtParam import ProteinAnalysis

# ==========================================
# 1. 底层序列元件与模板库
# ==========================================
TEMPLATES = {
    "linker": {
        "G4S_3": "GGGGSGGGGSGGGGS",
        "Rigid_18": "KPGSGKPGSGKPGSGKPG"
    },
    "ch1_hinge_front": {
        "IgG1_WT": "ASTKGPSVFPLAPSSKSTSGGTAALGCLVKDYFPEPVTVSWNSGALTSGVHTFPAVLQSSGLYSLSSVVTVPSSSLGTQTYICNVNHKPSNTKVDKKVEPKSC"
    },
    "fc_knob": {
        "LALA_GA": "EPKSSDKTHTCPPCPAPEAAGAPSVFLFPPKPKDTLMISRTPEVTCVVVDVSHEDPEVKFNWYVDGVEVHNAKTKPREEQYNSTYRVVSVLTVLHQDWLNGKEYKCKVSNKALPAPIEKTISKAKGQPREPQVYTLPPCRDELTKNQVSLWCLVKGFYPSDIAVEWESNGQPENNYKTTPPVLDSDGSFFLYSKLTVDKSRWQQGNVFSCSVMHEALHNHYTQKSLSLSPGK",
        "LALA_PG": "EPKSSDKTHTCPPCPAPEAAGGPSVFLFPPKPKDTLMISRTPEVTCVVVDVSHEDPEVKFNWYVDGVEVHNAKTKPREEQYNSTYRVVSVLTVLHQDWLNGKEYKCKVSNKALGAPIEKTISKAKGQPREPQVYTLPPCRDELTKNQVSLWCLVKGFYPSDIAVEWESNGQPENNYKTTPPVLDSDGSFFLYSKLTVDKSRWQQGNVFSCSVMHEALHNHYTQKSLSLSPGK"
    },
    "fc_hole_r": {
        "LALA_GA": "ASTKGPSVFPLAPSSKSTSGGTAALGCLVKDYFPEPVTVSWNSGALTSGVHTFPAVLQSSGLYSLSSVVTVPSSSLGTQTYICNVNHKPSNTKVDKKVEPKSCDKTHTCPPCPAPEAAGAPSVFLFPPKPKDTLMISRTPEVTCVVVDVSHEDPEVKFNWYVDGVEVHNAKTKPREEQYNSTYRVVSVLTVLHQDWLNGKEYKCKVSNKALPAPIEKTISKAKGQPREPQVCTLPPSRDELTKNQVSLSCAVKGFYPSDIAVEWESNGQPENNYKTTPPVLDSDGSFFLVSKLTVDKSRWQQGNVFSCSVMHEALHNRYTQKSLSLSPGK",
        "LALA_PG": "ASTKGPSVFPLAPSSKSTSGGTAALGCLVKDYFPEPVTVSWNSGALTSGVHTFPAVLQSSGLYSLSSVVTVPSSSLGTQTYICNVNHKPSNTKVDKKVEPKSCDKTHTCPPCPAPEAAGGPSVFLFPPKPKDTLMISRTPEVTCVVVDVSHEDPEVKFNWYVDGVEVHNAKTKPREEQYNSTYRVVSVLTVLHQDWLNGKEYKCKVSNKALGAPIEKTISKAKGQPREPQVCTLPPSRDELTKNQVSLSCAVKGFYPSDIAVEWESNGQPENNYKTTPPVLDSDGSFFLVSKLTVDKSRWQQGNVFSCSVMHEALHNRYTQKSLSLSPGK"
    },
    "light_constant": {
        "Kappa": "RTVAAPSVFIFPPSDEQLKSGTASVVCLLNNFYPREAKVQWKVDNALQSGNSQESVTEQDSKDSTYSLSSTLTLSKADYEKHKVYACEVTHQGLSSPVTKSFNRGEC",
        "Lambda": "GQPKANPTVTLFPPSSEELQANKATLVCLISDFYPGAVTVAWKADGSPVKAGVETTKPSKQSNNKYAASSYLSLTPEQWKSHRSYSCQVTHEGSTVEKTVAPTECS"
    }
}

COLOR_MAP = {
    'TAA_VH': RGBColor(31, 78, 121),
    'TAA_VL': RGBColor(0, 176, 240),
    'CD3_VH': RGBColor(192, 0, 0),
    'CD3_VL': RGBColor(255, 0, 0),
    'Linker': RGBColor(56, 87, 35),
    'Constant': RGBColor(127, 127, 127),
    'Fc': RGBColor(197, 90, 17)
}

# ==========================================
# 2. 核心功能函数 (解析、组装、计算)
# ==========================================
def parse_fasta(text):
    clones = {}
    current_name, current_chain = None, None
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('>'):
            header = line[1:].strip()
            parts = header.split('_')
            if len(parts) >= 2:
                current_chain = parts[-1].upper()
                current_name = '_'.join(parts[:-1])
                if current_name not in clones: clones[current_name] = {}
            else:
                current_name, current_chain = header, 'UNKNOWN'
        else:
            if current_name and current_chain in ['VH', 'VL']:
                seq = re.sub(r'[^a-zA-Z]', '', line).upper()
                if current_chain not in clones[current_name]: clones[current_name][current_chain] = ""
                clones[current_name][current_chain] += seq
    return {k: v for k, v in clones.items() if 'VH' in v and 'VL' in v}

def assemble_segments(taa_vh, taa_vl, cd3_vh, cd3_vl, format_type, fc_silence, linker, lc_type):
    chains = {}
    link_seq = TEMPLATES["linker"][linker]
    ch1_hinge = TEMPLATES["ch1_hinge_front"]["IgG1_WT"]
    fc_knob = TEMPLATES["fc_knob"][fc_silence]
    fc_hole = TEMPLATES["fc_hole_r"][fc_silence]
    ck = TEMPLATES["light_constant"][lc_type]
    
    if format_type.startswith("2+1"):
        chains['H1_scFv-Fc_Knob'] = [('TAA_VH', taa_vh), ('Constant', ch1_hinge), ('Linker', link_seq), ('CD3_VL', cd3_vl), ('Linker', link_seq), ('CD3_VH', cd3_vh), ('Fc', fc_knob)]
        chains['H2_Fab-Fc_Hole'] = [('TAA_VH', taa_vh), ('Fc', fc_hole)]
        chains['L1_Common_LC'] = [('TAA_VL', taa_vl), ('Constant', ck)]
    else:
        chains['H1_CD3_scFv-Fc_Knob'] = [('CD3_VL', cd3_vl), ('Linker', link_seq), ('CD3_VH', cd3_vh), ('Fc', fc_knob)]
        chains['H2_TAA_Fab-Fc_Hole'] = [('TAA_VH', taa_vh), ('Fc', fc_hole)]
        chains['L1_TAA_LC'] = [('TAA_VL', taa_vl), ('Constant', ck)]
    return chains

def analyze_developability(chains_dict, format_type):
    """计算单个双抗分子的理化参数"""
    raw_seqs = {name: "".join([seg[1] for seg in segments]) for name, segments in chains_dict.items()}
    chain_names = list(raw_seqs.keys())
    
    metrics = {}
    total_mw = 0
    
    try:
        # H1 分析
        h1_ana = ProteinAnalysis(raw_seqs[chain_names[0]])
        metrics['H1_pI'] = round(h1_ana.isoelectric_point(), 2)
        total_mw += h1_ana.molecular_weight()
        
        # H2 分析
        h2_ana = ProteinAnalysis(raw_seqs[chain_names[1]])
        metrics['H2_pI'] = round(h2_ana.isoelectric_point(), 2)
        total_mw += h2_ana.molecular_weight()
        
        # L1 分析
        l1_ana = ProteinAnalysis(raw_seqs[chain_names[2]])
        metrics['L1_pI'] = round(l1_ana.isoelectric_point(), 2)
        
        # 化学计量比计算完整分子量 (2+1 需要两条轻链，1+1 需要一条)
        if format_type.startswith("2+1"):
            total_mw += (l1_ana.molecular_weight() * 2)
        else:
            total_mw += l1_ana.molecular_weight()
            
        metrics['Total_MW(kDa)'] = round(total_mw / 1000, 2)
        metrics['ΔpI(H1-H2)'] = round(abs(metrics['H1_pI'] - metrics['H2_pI']), 2)
        
        # 整体稳定性预估 (取最差的那条链)
        metrics['Max_GRAVY'] = round(max(h1_ana.gravy(), h2_ana.gravy(), l1_ana.gravy()), 3)
        metrics['Max_Instability'] = round(max(h1_ana.instability_index(), h2_ana.instability_index(), l1_ana.instability_index()), 1)
        
    except Exception as e:
        metrics = {'Total_MW(kDa)': 0, 'H1_pI': 0, 'H2_pI': 0, 'L1_pI': 0, 'ΔpI(H1-H2)': 0, 'Max_GRAVY': 0, 'Max_Instability': 0}
        
    return metrics

def generate_word_document(combinations_data):
    doc = Document()
    doc.styles['Normal'].font.name = 'Courier New'
    doc.styles['Normal'].font.size = Pt(9)
    
    doc.add_heading('HTS Bispecific Antibody Sequence Report', 0)
    legend = doc.add_paragraph()
    legend.add_run("Color Legend: ").bold = True
    for key, color in COLOR_MAP.items():
        run = legend.add_run(f" [{key}] ")
        run.font.color.rgb = color
        run.bold = True
        
    for combo in combinations_data:
        doc.add_heading(combo['Name'], level=1)
        # 写入参数
        p_param = doc.add_paragraph()
        p_param.add_run(f"Pairing: TAA [{combo['TAA_Clone']}] × CD3 [{combo['CD3_Clone']}]\n").bold = True
        p_param.add_run(f"MW: {combo['Metrics']['Total_MW(kDa)']} kDa | ΔpI(H1-H2): {combo['Metrics']['ΔpI(H1-H2)']} | GRAVY: {combo['Metrics']['Max_GRAVY']}")
        
        for chain_name, segments in combo['Chains'].items():
            p = doc.add_paragraph()
            p.add_run(f">{combo['Name']}_{chain_name}\n").bold = True
            for seg_type, seq in segments:
                run = p.add_run(seq)
                if seg_type in COLOR_MAP: run.font.color.rgb = COLOR_MAP[seg_type]
                    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def convert_df_to_csv(df):
    return df.to_csv(index=False).encode('utf-8')

# ==========================================
# 3. Streamlit UI 布局
# ==========================================
st.set_page_config(page_title="TCE HTS Combinatorial Platform", page_icon="🔬", layout="wide")

st.title("🔬 工业级 TCE 双抗高通量组合与生成平台")
st.markdown("不仅能自动执行**交叉组合 (Cross-pairing)** 和**着色文档生成**，还能深度分析每一款双抗的 **CMC 成药性评估参数 (Developability)**。")

with st.sidebar:
    st.header("⚙️ 平台构建参数")
    st.markdown("---")
    base_name = st.text_input("📦 基础项目命名 (Base Name)", value="BMK")
    start_idx = st.number_input("🔢 起始编号", min_value=1, value=1)
    
    st.markdown("---")
    format_option = st.radio("📐 双抗构型 (Format)", ["2+1 构型 (TAA双臂 + CD3单臂)", "1+1 构型 (TAA单臂 + CD3单臂)"])
    fc_option = st.selectbox("🛡️ Fc 沉默底盘", ["LALA_GA", "LALA_PG"])
    linker_option = st.selectbox("🔗 scFv 接头类型", ["G4S_3", "Rigid_18"])
    lc_option = st.selectbox("⛓️ 恒定区轻链类型", ["Kappa", "Lambda"])
    
    st.markdown("---")
    st.info("**📈 参数解读指南**\n\n**ΔpI (H1-H2)**: 建议 >0.2 以利于离子交换层析纯化分离。\n\n**Max GRAVY**: 疏水性得分，正值越大越易发生聚集沉淀。\n\n**Total MW**: 质谱鉴定的绝对参考值。")

col1, col2 = st.columns(2)
with col1:
    st.subheader("🎯 TAA 克隆文库")
    taa_fasta = st.text_area("输入 TAA 序列", height=200, placeholder=">CloneA_VH\nEVQL...\n>CloneA_VL\nDIQM...")
with col2:
    st.subheader("⚔️ CD3 克隆文库")
    cd3_fasta = st.text_area("输入 CD3 序列", height=200, placeholder=">SP34_VH\nEVQL...\n>SP34_VL\nQTVV...")

if st.button("🚀 批量交叉组合、分析并生成文件", type="primary", use_container_width=True):
    if not taa_fasta or not cd3_fasta:
        st.warning("⚠️ 必须同时提供 TAA 和 CD3 的批量序列！")
    else:
        with st.spinner("正在执行矩阵组合、CMC参数计算及文件构建..."):
            taa_clones = parse_fasta(taa_fasta)
            cd3_clones = parse_fasta(cd3_fasta)
            
            if not taa_clones or not cd3_clones:
                st.error("❌ 无法解析序列，请确保严格采用 `>名称_VH` 和 `>名称_VL` 格式。")
            else:
                combinations = []
                counter = start_idx
                overview_data = []
                
                for taa_name, taa_seqs in taa_clones.items():
                    for cd3_name, cd3_seqs in cd3_clones.items():
                        combo_name = f"{base_name}-{counter:03d}"
                        counter += 1
                        
                        chains = assemble_segments(
                            taa_vh=taa_seqs['VH'], taa_vl=taa_seqs['VL'],
                            cd3_vh=cd3_seqs['VH'], cd3_vl=cd3_seqs['VL'],
                            format_type=format_option, fc_silence=fc_option, 
                            linker=linker_option, lc_type=lc_option
                        )
                        
                        metrics = analyze_developability(chains, format_option)
                        
                        combinations.append({
                            'Name': combo_name,
                            'TAA_Clone': taa_name,
                            'CD3_Clone': cd3_name,
                            'Chains': chains,
                            'Metrics': metrics
                        })
                        
                        overview_data.append({
                            "分子编号": combo_name,
                            "TAA 来源": taa_name,
                            "CD3 来源": cd3_name,
                            "Total MW (kDa)": metrics['Total_MW(kDa)'],
                            "ΔpI (H1-H2)": metrics['ΔpI(H1-H2)'],
                            "Max GRAVY": metrics['Max_GRAVY'],
                            "Instability": metrics['Max_Instability'],
                            "H1_pI": metrics['H1_pI'],
                            "H2_pI": metrics['H2_pI']
                        })
                
                st.success(f"✅ 成功生成并分析了 {len(combinations)} 种双抗组合！")
                
                # 展现数据总览大表
                st.subheader("📊 CMC 成药性评估总表 (Developability Overview)")
                df_overview = pd.DataFrame(overview_data)
                
                # 使用 Streamlit Dataframe 渲染，带有高亮和排序功能
                st.dataframe(
                    df_overview.style.background_gradient(subset=['ΔpI (H1-H2)'], cmap='Greens')
                                     .background_gradient(subset=['Max GRAVY'], cmap='OrRd'),
                    use_container_width=True
                )
                
                # 提供下载按钮区
                st.subheader("💾 文件下载区")
                col_btn1, col_btn2 = st.columns(2)
                
                with col_btn1:
                    word_file = generate_word_document(combinations)
                    st.download_button(
                        label="📄 下载可视化着色序列报告 (Word)",
                        data=word_file,
                        file_name=f"{base_name}_HTS_Library_Sequences.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                with col_btn2:
                    csv_file = convert_df_to_csv(df_overview)
                    st.download_button(
                        label="📊 下载成药性评估参数表 (CSV / Excel可用)",
                        data=csv_file,
                        file_name=f"{base_name}_Developability_Report.csv",
                        mime="text/csv",
                        use_container_width=True
                    )
